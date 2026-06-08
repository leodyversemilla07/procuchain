#!/usr/bin/env python3
"""Generate 2-column US Letter DOCX of ProcuChain source code.

Excludes generated code:
  - resources/js/actions/       (auto-generated TS route mirrors)
  - resources/js/routes/        (Wayfinder-generated route definitions — 66 files)
  - resources/js/types/vite-env.d.ts
  - public/build/, tests, vendor, node_modules, .d.ts

Strips all comments from code (PHP, JS/TS single-line and multi-line).
No title page, no overview, no file index — just source code.
"""

import os
import re
import time
from pathlib import Path

PROJECT_ROOT = Path("/home/ubuntu/workspace/procuchain")
OUTPUT_DOCX = PROJECT_ROOT / "PROCUCHAIN_SOURCE_DOCS.docx"

EXCLUDE_DIRS = {
    "node_modules", "vendor", ".git", "storage", "bootstrap",
    "public", "tests", "__pycache__", ".elasticbeanstalk",
}

# Generated code patterns to exclude
EXCLUDE_PATTERNS = [
    "resources/js/actions/",       # auto-generated controller mirrors
    "resources/js/routes/",        # Wayfinder-generated route definitions
    "vite-env.d.ts",              # Vite boilerplate
]

EXCLUDE_EXTENSIONS = {".d.ts"}
MAX_LINES_PER_FILE = 300
CHUNK_SIZE = 80

SECTIONS = [
    ("app/Console/Commands", "Console Commands"),
    ("app/Contracts", "Contracts (Interfaces)"),
    ("app/DataTransferObjects", "Data Transfer Objects"),
    ("app/Enums", "Enums"),
    ("app/Events", "Events"),
    ("app/Exceptions", "Exceptions"),
    ("app/Http/Controllers", "HTTP Controllers"),
    ("app/Http/Middleware", "HTTP Middleware"),
    ("app/Http/Requests", "HTTP Requests (Validation)"),
    ("app/Http/Responses", "HTTP Responses"),
    ("app/Jobs", "Jobs"),
    ("app/Libraries", "Libraries"),
    ("app/Listeners", "Listeners"),
    ("app/Mail", "Mail (Mailables)"),
    ("app/Models", "Eloquent Models"),
    ("app/Notifications", "Notifications"),
    ("app/Policies", "Authorization Policies"),
    ("app/Providers", "Service Providers"),
    ("app/Repositories", "Repositories"),
    ("app/Services", "Services"),
    ("database/migrations", "Database Migrations"),
    ("database/seeders", "Database Seeders"),
    ("database/factories", "Database Factories"),
    ("routes", "Routes"),
    ("config", "Configuration"),
    ("resources/views", "Blade Templates"),
    ("resources/js/pages", "React Pages"),
    ("resources/js/components", "React Components"),
    ("resources/js/hooks", "React Hooks"),
    ("resources/js/types", "TypeScript Types"),
    ("resources/js/lib", "TypeScript Libraries"),
    # routes/ excluded — Wayfinder generated
    ("resources/js", "Frontend Root (app.tsx)"),
]

VALID_EXT = {".php", ".ts", ".tsx"}


def should_include(rel_path: str) -> bool:
    parts = rel_path.split("/")
    for part in parts:
        if part in EXCLUDE_DIRS:
            return False
    for pattern in EXCLUDE_PATTERNS:
        if pattern in rel_path:
            return False
    for ext in EXCLUDE_EXTENSIONS:
        if rel_path.endswith(ext):
            return False
    _, file_ext = os.path.splitext(rel_path)
    return file_ext in VALID_EXT


def get_files_for_section(section_dir: str) -> list[str]:
    full_dir = PROJECT_ROOT / section_dir
    if not full_dir.exists():
        return []
    files = []
    for root, dirs, filenames in os.walk(full_dir):
        dirs[:] = [d for d in dirs if d not in EXCLUDE_DIRS]
        for fname in sorted(filenames):
            fpath = os.path.join(root, fname)
            rel_path = os.path.relpath(fpath, PROJECT_ROOT)
            if should_include(rel_path):
                files.append(rel_path)
    return sorted(files)


def strip_comments(code: str, lang: str) -> str:
    """Remove all comments from source code."""
    if lang == "php":
        # Remove // comments (but not URLs inside strings — keep it simple)
        code = re.sub(r'(?<!:)//.*$', '', code, flags=re.MULTILINE)
        # Remove # comments
        code = re.sub(r'#.*$', '', code, flags=re.MULTILINE)
        # Remove /* ... */ block comments
        code = re.sub(r'/\*.*?\*/', '', code, flags=re.DOTALL)
        # Remove /** ... */ docblocks
        code = re.sub(r'/\*\*.*?\*/', '', code, flags=re.DOTALL)
    elif lang in ("ts", "tsx"):
        # Remove // comments
        code = re.sub(r'(?<!:)//.*$', '', code, flags=re.MULTILINE)
        # Remove /* ... */ block comments (includes JSDoc)
        code = re.sub(r'/\*.*?\*/', '', code, flags=re.DOTALL)
    # Collapse multiple blank lines to max 2
    code = re.sub(r'\n{3,}', '\n\n', code)
    return code.strip()


def read_lines(rel_path: str, lang: str) -> list[str]:
    full_path = PROJECT_ROOT / rel_path
    try:
        with open(full_path, "r", encoding="utf-8", errors="replace") as f:
            raw = f.read()
    except (OSError, IOError):
        return ["*File not readable*"]

    # Strip comments
    clean = strip_comments(raw, lang)

    lines = clean.split("\n")
    if len(lines) > MAX_LINES_PER_FILE:
        lines = lines[:MAX_LINES_PER_FILE]
        lines.append(f"// ... truncated (showing {MAX_LINES_PER_FILE} of {len(lines)} lines)")

    # Filter out empty-only lines at start/end, but keep internal blanks
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()

    return lines


def detect_language(rel_path: str) -> str:
    if rel_path.endswith(".tsx"):
        return "tsx"
    if rel_path.endswith(".ts"):
        return "ts"
    if rel_path.endswith(".php"):
        return "php"
    return "text"


def count_file_lines(rel_path: str) -> int:
    try:
        with open(PROJECT_ROOT / rel_path, "r", encoding="utf-8", errors="replace") as f:
            return sum(1 for _ in f)
    except:
        return 0


def generate_docx():
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    t0 = time.time()
    doc = Document()

    # ── Page setup: US Letter ──
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)

    # ── 2-column layout ──
    sectPr = section._sectPr
    cols = parse_xml(
        '<w:cols xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:num="2" w:space="360" w:sep="1"/>'
    )
    sectPr.append(cols)

    # ── Styles ──
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(7)
    style_normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style_normal.paragraph_format.space_before = Pt(0)
    style_normal.paragraph_format.space_after = Pt(1)
    style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    h1 = doc.styles['Heading 1']
    h1.font.size = Pt(14)
    h1.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    h1.font.bold = True
    h1.font.name = 'Calibri'
    h1.paragraph_format.space_before = Pt(8)
    h1.paragraph_format.space_after = Pt(4)

    h2 = doc.styles['Heading 2']
    h2.font.size = Pt(9)
    h2.font.color.rgb = RGBColor(0x20, 0x4a, 0x87)
    h2.font.bold = True
    h2.font.name = 'Calibri'
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(2)

    h3 = doc.styles['Heading 3']
    h3.font.size = Pt(7)
    h3.font.color.rgb = RGBColor(0x4e, 0x9a, 0x06)
    h3.font.bold = True
    h3.font.name = 'Consolas'
    h3.paragraph_format.space_before = Pt(4)
    h3.paragraph_format.space_after = Pt(1)

    code_style = doc.styles.add_style('CodeBlock', 1)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(5)
    code_style.font.color.rgb = RGBColor(0x2e, 0x2e, 0x2e)
    code_style.paragraph_format.space_before = Pt(0)
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    code_style.element.append(shading)

    # ── Header ──
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "ProCuChain Source Code — RA 12009"
    hp.alignment = 1
    for run in hp.runs:
        run.font.size = Pt(6)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ═══════════════════════════════════════════
    # SOURCE CODE ONLY — no title, no index
    # ═══════════════════════════════════════════
    file_count = 0
    total_files = 0
    total_lines = 0
    para_count = 0

    for section_dir, section_title in SECTIONS:
        files = get_files_for_section(section_dir)
        if not files:
            continue

        doc.add_heading(section_title, level=2)
        total_files += len(files)

        for rel_path in files:
            file_count += 1
            lang = detect_language(rel_path)
            lines = read_lines(rel_path, lang)
            total_lines += len(lines)

            # File heading
            doc.add_heading(rel_path, level=3)

            # Code in chunked batches
            for i in range(0, len(lines), CHUNK_SIZE):
                chunk = lines[i:i + CHUNK_SIZE]
                p = doc.add_paragraph(style='CodeBlock')
                run = p.add_run("\n".join(chunk))
                run.font.name = 'Consolas'
                run.font.size = Pt(5)
                para_count += 1

            # Thin separator
            sep = doc.add_paragraph()
            sep_run = sep.add_run("─" * 60)
            sep_run.font.size = Pt(4)
            sep_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

            if file_count % 20 == 0:
                elapsed = time.time() - t0
                print(f"  [{elapsed:.0f}s] {file_count} files, {para_count} paragraphs")

    # ── Save ──
    print(f"\nSaving DOCX...")
    doc.save(str(OUTPUT_DOCX))
    size = OUTPUT_DOCX.stat().st_size
    elapsed = time.time() - t0
    print(f"\n{'=' * 60}")
    print(f" ProcuChain Source Code (DOCX) — comments stripped")
    print(f" Files: {total_files} (excluded: actions/ + routes/ Wayfinder)")
    print(f" Lines: {total_lines:,}")
    print(f" Size: {size / 1024:.0f} KB")
    print(f" Time: {elapsed:.0f}s")
    print(f" Output: {OUTPUT_DOCX}")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    generate_docx()
