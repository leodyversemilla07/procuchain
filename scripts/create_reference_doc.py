#!/usr/bin/env python3
"""Create a Word reference doc with 2-column layout and 5.5pt monospace code fonts."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── Page setup: A4, narrow margins ──
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(1.3)
section.right_margin = Cm(1.3)

# ── 2-column layout via XML ──
sectPr = section._sectPr
cols = parse_xml(
    '<w:cols xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'w:num="2" w:space="360" w:sep="1"/>'
)
sectPr.append(cols)

# ── Default style: 7pt ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(7)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(1)
pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

# ── Heading 1 ──
h1 = doc.styles['Heading 1']
h1.font.size = Pt(14)
h1.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
h1.font.bold = True
h1.font.name = 'Calibri'
h1.paragraph_format.space_before = Pt(8)
h1.paragraph_format.space_after = Pt(4)

# ── Heading 2 ──
h2 = doc.styles['Heading 2']
h2.font.size = Pt(9)
h2.font.color.rgb = RGBColor(0x20, 0x4a, 0x87)
h2.font.bold = True
h2.font.name = 'Calibri'
h2.paragraph_format.space_before = Pt(6)
h2.paragraph_format.space_after = Pt(2)

# ── Heading 3 ──
h3 = doc.styles['Heading 3']
h3.font.size = Pt(7)
h3.font.color.rgb = RGBColor(0x4e, 0x9a, 0x06)
h3.font.bold = True
h3.font.name = 'Consolas'
h3.paragraph_format.space_before = Pt(4)
h3.paragraph_format.space_after = Pt(1)

# ── Code Block style (pandoc uses "SourceCode" for fenced code blocks in docx) ──
# Pandoc looks for styles named "Source Code" in the reference doc
code_style = doc.styles.add_style('Source Code', 1)  # 1 = paragraph type
code_style.font.name = 'Consolas'
code_style.font.size = Pt(5)
code_style.font.color.rgb = RGBColor(0x2e, 0x2e, 0x2e)
code_style.paragraph_format.space_before = Pt(0)
code_style.paragraph_format.space_after = Pt(0)
code_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
# Shading for code blocks (light gray)
shading = parse_xml(
    f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>'
)
code_style.element.append(shading)

# Also add "Code" style as fallback
code2 = doc.styles.add_style('Code', 1)
code2.font.name = 'Consolas'
code2.font.size = Pt(5)
code2.font.color.rgb = RGBColor(0x2e, 0x2e, 0x2e)
code2.paragraph_format.space_before = Pt(0)
code2.paragraph_format.space_after = Pt(0)
code2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
shading2 = parse_xml(
    f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>'
)
code2.element.append(shading2)

# ── Header ──
header = section.header
hp = header.paragraphs[0]
run = hp.add_run("ProCuChain Source Code — RA 12009")
run.font.size = Pt(6)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
hp.alignment = 2

# ── Footer with page number ──
footer = section.footer
fp = footer.paragraphs[0]
run = fp.add_run()
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run._r.append(fldChar1)
run2 = fp.add_run()
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
run2._r.append(instrText)
run3 = fp.add_run()
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run3._r.append(fldChar2)
fp.alignment = 2

doc.add_paragraph("Reference document.")

out = "/home/ubuntu/workspace/procuchain/reference_doc.docx"
doc.save(out)
print(f"Saved: {out}")
